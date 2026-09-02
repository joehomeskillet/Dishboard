#!/usr/bin/env python3
"""Erzeugt das Word-SDD aus der Markdown-Quelle und den Referenzbildern."""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

DARK = '173B3F'
TEAL = '2F6768'
TEAL_LIGHT = 'E6F0EE'
MAGENTA = 'A52A62'
INK = '172A2D'
MUTED = '52666A'
GRID = 'C8D5D3'
PANEL = 'F3F7F6'
WHITE = 'FFFFFF'
FONT = 'Aptos'
MONO = 'Liberation Mono'


def set_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None, name: str = FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn('w:shd'))
    if node is None:
        node = OxmlElement('w:shd')
        tc_pr.append(node)
    node.set(qn('w:fill'), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn('w:shd'))
    if node is None:
        node = OxmlElement('w:shd')
        p_pr.append(node)
    node.set(qn('w:fill'), fill)


def margins(cell, value: int = 95) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge in ('top', 'start', 'bottom', 'end'):
        node = tc_mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement('w:tblHeader')
    node.set(qn('w:val'), 'true')
    tr_pr.append(node)


def dont_split(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))


def set_alt(shape, text: str) -> None:
    clean = re.sub(r'\s+', ' ', text).strip()
    shape._inline.docPr.set('title', clean)
    shape._inline.docPr.set('descr', clean)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run('Seite ')
    set_font(prefix, size=8, color=MUTED)
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' PAGE '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def inline(paragraph, text: str, *, size: float | None = None, color: str | None = None) -> None:
    pattern = re.compile(r'(\*\*.+?\*\*|`.+?`)')
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith('**'):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=(size or 9.5) - 0.3, color=DARK, name=MONO)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_font(run, size=size, color=color)


def setup_styles(doc: Document) -> None:
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [('Title', 28, DARK), ('Subtitle', 14, TEAL)]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    for level, size, color in [(1, 18, DARK), (2, 13, TEAL), (3, 10.5, MAGENTA)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if level > 1 else 4)
        style.paragraph_format.space_after = Pt(6)

    for list_name in ('List Bullet', 'List Number'):
        style = doc.styles[list_name]
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(9.2)
        style.font.color.rgb = RGBColor.from_string(INK)


def setup_section(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.55)
        section.bottom_margin = Cm(1.45)
        section.left_margin = Cm(1.65)
        section.right_margin = Cm(1.65)
        section.header_distance = Cm(0.65)
        section.footer_distance = Cm(0.65)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header.add_run('KLINIK SÜDHANG  |  MENÜPLANUNG  |  ENTWURF')
        set_font(run, size=7.5, bold=True, color=TEAL)
        footer = section.footer.paragraphs[0]
        add_page_number(footer)


def add_cover(doc: Document, root: Path) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('KLINIK SÜDHANG')
    set_font(r, size=12, bold=True, color=MAGENTA)

    title = doc.add_paragraph(style='Title')
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(5)
    title.add_run('Menüplanung')

    subtitle = doc.add_paragraph(style='Subtitle')
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.add_run('Getrennte Patienten- und Cafeteria-Publikation')

    status = doc.add_table(rows=2, cols=3)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    status.autofit = True
    repeat_header(status.rows[0])
    for row in status.rows:
        dont_split(row)
    values = [('STATUS', 'Entwurf / intern geprüft'), ('STAND', '1. September 2026'), ('ABNAHME', 'fachlich offen')]
    for col, (label, value) in enumerate(values):
        c1 = status.cell(0, col)
        c2 = status.cell(1, col)
        shade(c1, TEAL_LIGHT)
        shade(c2, PANEL)
        margins(c1, 80)
        margins(c2, 95)
        for c in (c1, c2):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline(c1.paragraphs[0], label, size=7.5, color=MUTED)
        inline(c2.paragraphs[0], value, size=9.5)

    doc.add_paragraph()
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_together = True
    images = [
        ('design/screenshots/signage-cafeteria-tag-1920x1080.png', 'Cafeteria-Player mit zwei Menüarten und zwei Kostenansätzen'),
        ('design/screenshots/signage-patienten-tag-1920x1080.png', 'Patienten-Player mit Mittag und Abend sowie je zwei Menüarten'),
    ]
    for idx, (relative, alt) in enumerate(images):
        if idx:
            spacer = image_paragraph.add_run('  ')
            set_font(spacer, size=3)
        shape = image_paragraph.add_run().add_picture(str(root / relative), width=Cm(8.1))
        set_alt(shape, alt)

    message = doc.add_paragraph()
    message.alignment = WD_ALIGN_PARAGRAPH.CENTER
    message.paragraph_format.left_indent = Cm(0.15)
    message.paragraph_format.right_indent = Cm(0.15)
    message.paragraph_format.space_before = Pt(5)
    message.paragraph_format.space_after = Pt(5)
    shade_paragraph(message, DARK)
    run = message.add_run('Zwei Profile. Zwei Raster. Vier feste Player. Keine Kosteninformation im Patientenkanal.')
    set_font(run, size=11, bold=True, color=WHITE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    inline(p, 'Flask · PostgreSQL · Microsoft Entra ID · Docker Compose · Digital Signage', size=9, color=MUTED)
    doc.add_page_break()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
        if all(re.fullmatch(r':?-{3,}:?', cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        dont_split(row)
        if row_index == 0:
            repeat_header(row)
        for col_index in range(columns):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell, 70)
            if row_index == 0:
                shade(cell, DARK)
            elif row_index % 2 == 0:
                shade(cell, PANEL)
            text = values[col_index] if col_index < len(values) else ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            inline(p, text, size=7.8 if columns >= 5 else 8.3, color=WHITE if row_index == 0 else INK)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def resolve_image(source: Path, value: str) -> Path:
    path = Path(value)
    if path.is_absolute():
        return path
    return (source.parent / path).resolve()


def build(source: Path, output: Path) -> None:
    root = source.resolve().parents[1]
    doc = Document()
    setup_styles(doc)
    setup_section(doc)
    add_cover(doc, root)

    raw_lines = source.read_text(encoding='utf-8').splitlines()
    # Remove YAML front matter.
    if raw_lines and raw_lines[0].strip() == '---':
        try:
            end = raw_lines.index('---', 1)
            raw_lines = raw_lines[end + 1:]
        except ValueError:
            pass

    index = 0
    first_h1 = True
    while index < len(raw_lines):
        line = raw_lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith('```'):
            code: list[str] = []
            index += 1
            while index < len(raw_lines) and not raw_lines[index].strip().startswith('```'):
                code.append(raw_lines[index])
                index += 1
            index += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.2)
            p.paragraph_format.right_indent = Cm(0.2)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_together = True
            shade_paragraph(p, PANEL)
            run = p.add_run('\n'.join(code))
            set_font(run, size=7.1, color=DARK, name=MONO)
            continue

        image_match = re.fullmatch(r'!\[(.*?)\]\((.*?)\)', stripped)
        if image_match:
            alt, value = image_match.groups()
            image_path = resolve_image(source, value)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_together = True
            if image_path.is_file():
                with Image.open(image_path) as source_image:
                    aspect = source_image.height / source_image.width
                if aspect > 1.2:
                    shape = p.add_run().add_picture(str(image_path), height=Cm(20.6))
                else:
                    shape = p.add_run().add_picture(str(image_path), width=Cm(16.5))
                set_alt(shape, alt)
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.keep_with_next = False
                r = caption.add_run(alt)
                set_font(r, size=7.8, color=MUTED)
            else:
                inline(p, f'[Bild fehlt: {value}]', color=MAGENTA)
            index += 1
            continue

        heading = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not first_h1:
                doc.add_page_break()
            paragraph = doc.add_heading(level=level)
            inline(paragraph, text, size={1: 18, 2: 13, 3: 10.5}[level], color={1: DARK, 2: TEAL, 3: MAGENTA}[level])
            if level == 1:
                first_h1 = False
            index += 1
            continue

        if stripped.startswith('|'):
            table_lines: list[str] = []
            while index < len(raw_lines) and raw_lines[index].strip().startswith('|'):
                table_lines.append(raw_lines[index].strip())
                index += 1
            add_table(doc, parse_table(table_lines))
            continue

        bullet = re.match(r'^-\s+(.+)$', stripped)
        numbered = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if bullet:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            inline(p, bullet.group(1))
            index += 1
            continue
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            p.paragraph_format.space_after = Pt(2)
            inline(p, f'{numbered.group(1)}. {numbered.group(2)}')
            index += 1
            continue

        # Normal paragraph: join wrapped Markdown lines until next block.
        parts = [stripped]
        index += 1
        while index < len(raw_lines):
            nxt = raw_lines[index].strip()
            if not nxt or nxt.startswith(('#', '|', '```', '![', '- ')) or re.match(r'^\d+\.\s+', nxt):
                break
            parts.append(nxt)
            index += 1
        p = doc.add_paragraph()
        inline(p, ' '.join(parts))

    add_header_footer(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = 'Menüplanung Klinik Südhang - Software Design Document'
    doc.core_properties.subject = 'Getrennte Patienten- und Cafeteria-Publikation'
    doc.core_properties.author = 'Klinik Südhang'
    doc.core_properties.keywords = 'Cafeteria, Patienten, Speiseplan, Digital Signage, Flask, PostgreSQL'
    doc.save(output)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    default_root = Path(__file__).resolve().parents[1]
    parser.add_argument('--source', type=Path, default=default_root / 'docs/SDD_Klinik_Suedhang_Cafeteria_v3.0.md')
    parser.add_argument('--output', type=Path, default=default_root / 'docs/SDD_Klinik_Suedhang_Cafeteria_v3.0.docx')
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())
